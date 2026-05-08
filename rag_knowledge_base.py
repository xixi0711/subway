"""
RAG知识库模块 - 支持文档检索增强
基于本地向量存储
"""

import os
import json
import math
from typing import List, Dict, Any, Optional
from dataclasses import dataclass


@dataclass
class Document:
    """文档结构"""
    content: str
    metadata: Dict[str, Any]


class LocalVectorStore:
    """本地向量存储实现"""
    
    def __init__(self, persist_file: str = "./vector_store.json"):
        self.persist_file = persist_file
        self.documents = []
        self.embeddings = []
        self.metadata = []
        self.vocabulary = set()
        self._load_from_disk()
    
    def _load_from_disk(self):
        """从磁盘加载向量存储"""
        try:
            if os.path.exists(self.persist_file):
                with open(self.persist_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.documents = data.get('documents', [])
                    self.embeddings = data.get('embeddings', [])
                    self.metadata = data.get('metadata', [])
                    self.vocabulary = set(data.get('vocabulary', []))
        except Exception as e:
            print(f"加载向量存储失败: {e}")
    
    def _save_to_disk(self):
        """保存向量存储到磁盘"""
        try:
            data = {
                'documents': self.documents,
                'embeddings': self.embeddings,
                'metadata': self.metadata,
                'vocabulary': list(self.vocabulary)
            }
            with open(self.persist_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"保存向量存储失败: {e}")
    
    def _tokenize(self, text: str) -> List[str]:
        """简单分词"""
        import re
        # 简单的中文分词，基于字符
        return list(re.sub(r'\s+', '', text))
    
    def _get_embedding(self, text: str) -> List[float]:
        """生成文本的向量表示"""
        tokens = self._tokenize(text)
        # 构建词频向量
        vector = []
        for word in self.vocabulary:
            vector.append(tokens.count(word) / len(tokens) if tokens else 0)
        # 如果词汇表为空，返回空向量
        if not vector:
            return []
        # 归一化
        norm = math.sqrt(sum(x*x for x in vector))
        if norm > 0:
            vector = [x/norm for x in vector]
        return vector
    
    def add_texts(self, texts: List[str], metadatas: List[Dict[str, Any]]):
        """添加文本到向量存储"""
        for text, metadata in zip(texts, metadatas):
            # 更新词汇表
            tokens = self._tokenize(text)
            self.vocabulary.update(tokens)
        
        # 重新计算所有向量
        self.embeddings = []
        for text in self.documents + texts:
            self.embeddings.append(self._get_embedding(text))
        
        # 添加新文档
        self.documents.extend(texts)
        self.metadata.extend(metadatas)
        
        # 保存到磁盘
        self._save_to_disk()
    
    def similarity_search_with_score(self, query: str, k: int = 3) -> List[tuple]:
        """相似度搜索"""
        if not self.documents:
            return []
        
        # 生成查询向量
        query_embedding = self._get_embedding(query)
        if not query_embedding:
            return []
        
        # 计算相似度
        scores = []
        for i, embedding in enumerate(self.embeddings):
            if not embedding:
                continue
            # 余弦相似度
            dot_product = sum(a*b for a, b in zip(query_embedding, embedding))
            scores.append((i, dot_product))
        
        # 排序
        scores.sort(key=lambda x: x[1], reverse=True)
        
        # 返回结果
        results = []
        for idx, score in scores[:k]:
            class Doc:
                def __init__(self, content, metadata):
                    self.page_content = content
                    self.metadata = metadata
            results.append((Doc(self.documents[idx], self.metadata[idx]), score))
        
        return results
    
    def persist(self):
        """持久化到磁盘"""
        self._save_to_disk()
    
    def delete_collection(self):
        """删除集合"""
        self.documents = []
        self.embeddings = []
        self.metadata = []
        self.vocabulary = set()
        self._save_to_disk()


class SimpleRAG:
    """简化的RAG实现，支持基础文档检索"""

    def __init__(self, persist_directory: str = "./vector_store"):
        self.persist_directory = persist_directory
        self.documents = []
        self.embeddings = []
        self.metadata = []
        self.vectorstore = None
        self._initialized = False

    def _init_vectorstore(self):
        """初始化向量数据库"""
        if not self.vectorstore:
            try:
                # 创建持久化目录
                os.makedirs(self.persist_directory, exist_ok=True)
                persist_file = os.path.join(self.persist_directory, "vector_store.json")
                self.vectorstore = LocalVectorStore(persist_file=persist_file)
                self._initialized = True
                print("本地向量存储初始化成功")
            except Exception as e:
                print(f"初始化向量数据库失败: {e}")
                self.vectorstore = None

    def add_documents(self, documents: List[Document]) -> bool:
        """添加文档到知识库"""
        try:
            self._init_vectorstore()

            if self.vectorstore is None:
                for doc in documents:
                    self.documents.append(doc.content)
                    self.metadata.append(doc.metadata)
                return True

            texts = [doc.content for doc in documents]
            metas = [doc.metadata for doc in documents]

            self.vectorstore.add_texts(texts=texts, metadatas=metas)
            self.vectorstore.persist()

            return True
        except Exception as e:
            print(f"添加文档失败: {e}")
            for doc in documents:
                self.documents.append(doc.content)
                self.metadata.append(doc.metadata)
            return True

    def retrieve(self, query: str, k: int = 3) -> List[Dict[str, Any]]:
        """检索相关文档"""
        try:
            self._init_vectorstore()

            if self.vectorstore is None:
                return self._keyword_search(query, k)

            results = self.vectorstore.similarity_search_with_score(query, k=k)

            retrieved = []
            for doc, score in results:
                retrieved.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "score": float(score)
                })

            return retrieved
        except Exception as e:
            print(f"检索失败: {e}")
            return self._keyword_search(query, k)

    def _keyword_search(self, query: str, k: int = 3) -> List[Dict[str, Any]]:
        """关键词搜索作为fallback"""
        if not self.documents:
            return []

        query_words = set(query)
        scores = []

        for i, doc in enumerate(self.documents):
            doc_words = set(doc)
            overlap = len(query_words & doc_words)
            if overlap > 0:
                scores.append((i, overlap / len(query_words)))

        scores.sort(key=lambda x: x[1], reverse=True)

        results = []
        for idx, score in scores[:k]:
            results.append({
                "content": self.documents[idx],
                "metadata": self.metadata[idx] if idx < len(self.metadata) else {},
                "score": score
            })

        return results

    def load_documents_from_file(self, file_path: str) -> List[Document]:
        """从文件加载文档"""
        documents = []
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    if isinstance(data, list):
                        for item in data:
                            if isinstance(item, dict):
                                content = item.get('content', str(item))
                                metadata = {k: v for k, v in item.items() if k != 'content'}
                            else:
                                content = str(item)
                                metadata = {}
                            documents.append(Document(content=content, metadata=metadata))
                    elif isinstance(data, dict):
                        content = data.get('content', str(data))
                        metadata = {k: v for k, v in data.items() if k != 'content'}
                        documents.append(Document(content=content, metadata=metadata))

            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    documents.append(Document(
                        content=content,
                        metadata={"source": file_path, "type": "text"}
                    ))

            elif file_ext == '.pdf':
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        for i, page in enumerate(reader.pages):
                            text = page.extract_text()
                            if text.strip():
                                documents.append(Document(
                                    content=text,
                                    metadata={"source": file_path, "page": i + 1, "type": "pdf"}
                                ))
                except ImportError:
                    print("警告: PyPDF2未安装，无法解析PDF")

            elif file_ext == '.docx':
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                    if content.strip():
                        documents.append(Document(
                            content=content,
                            metadata={"source": file_path, "type": "docx"}
                        ))
                except ImportError:
                    print("警告: python-docx未安装，无法解析DOCX")

        except Exception as e:
            print(f"加载文件 {file_path} 失败: {e}")

        return documents

    def build_from_directory(self, directory: str) -> int:
        """从目录批量加载文档"""
        if not os.path.exists(directory):
            print(f"目录不存在: {directory}")
            return 0

        count = 0
        for filename in os.listdir(directory):
            file_path = os.path.join(directory, filename)
            if os.path.isfile(file_path):
                docs = self.load_documents_from_file(file_path)
                if docs:
                    self.add_documents(docs)
                    count += len(docs)
                    print(f"已加载: {filename} ({len(docs)} 个文档)")

        return count

    def clear(self) -> None:
        """清空知识库"""
        self.documents.clear()
        self.metadata.clear()
        self.embeddings.clear()

        try:
            if self.vectorstore:
                self.vectorstore.delete_collection()
                self.vectorstore = None
        except:
            pass

        self._initialized = False


rag_system = SimpleRAG(persist_directory="./vector_store")


def add_knowledge(content: str, metadata: Dict[str, Any] = None) -> bool:
    """添加单条知识"""
    if metadata is None:
        metadata = {}
    doc = Document(content=content, metadata=metadata)
    return rag_system.add_documents([doc])


def retrieve_knowledge(query: str, k: int = 3) -> List[Dict[str, Any]]:
    """检索知识"""
    return rag_system.retrieve(query, k=k)


def build_knowledge_base(directory: str = "./knowledge_base") -> int:
    """从目录构建知识库"""
    return rag_system.build_from_directory(directory)


def get_retrieval_context(query: str, k: int = 3) -> str:
    """获取检索上下文，用于增强prompt"""
    results = retrieve_knowledge(query, k=k)

    if not results:
        return ""

    context = "相关知识参考：\n"
    for i, result in enumerate(results, 1):
        context += f"\n【文档{i}】(相关度:{result['score']:.2f})\n"
        context += result['content'][:500]
        if len(result['content']) > 500:
            context += "..."
        context += "\n"

    return context
